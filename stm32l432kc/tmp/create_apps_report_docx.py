from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/APPS_CAN_사전검증_보고서.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
FONT = "Malgun Gothic"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name=FONT, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_para_format(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        style_size, color, before, after = 13, BLUE, 12, 6
    else:
        style_size, color, before, after = 12, DARK_BLUE, 8, 4
    set_para_format(p, before, after, 1.10)
    run = p.add_run(text)
    set_run_font(run, size=style_size, bold=True, color=color)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, 0, 6, 1.10)
    for idx, part in enumerate(text.split("`")):
        run = p.add_run(part)
        if idx % 2 == 1:
            set_run_font(run, "Consolas", 10, False, DARK_BLUE)
        else:
            set_run_font(run, size=11, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, 0, 8, 1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_para_format(p, 0, 8, 1.167)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_cell_margins(table, 120, 160, 120, 160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para_format(p, 0, 0, 1.10)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, "Consolas", 9.5, False, RGBColor(30, 55, 80))
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        set_para_format(p, 0, 0, 1.10)
        r = p.add_run(text)
        set_run_font(r, size=10.5, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            set_para_format(p, 0, 0, 1.10)
            r = p.add_run(str(text))
            if str(text).startswith("0x") or "CAN_" in str(text) or str(text).startswith("apps "):
                set_run_font(r, "Consolas", 9.5, False, INK)
            else:
                set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph()
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        st = styles[style_name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(11)
        st.paragraph_format.left_indent = Inches(0.5)
        st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(8)
        st.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("APPS-CAN 사전 검증 보고서")
    set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_format(title, 0, 4, 1.10)
    r = title.add_run("APPS 기반 CAN 모터 제어 사전 검증 보고서")
    set_run_font(r, size=22, bold=True, color=RGBColor(11, 37, 69))

    subtitle = doc.add_paragraph()
    set_para_format(subtitle, 0, 12, 1.10)
    r = subtitle.add_run("STM32L432KC 보드 단독 검증 및 BAMOCAR 연동 준비")
    set_run_font(r, size=11, color=MUTED)

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["작성 목적", "APPS 타당성 검사와 CAN torque command 생성 구조의 사전 검증"],
            ["검증 대상", "STM32L432KC CLI, CAN driver, APPS plausibility logic"],
            ["현재 범위", "실제 APPS 센서와 BAMOCAR 연결 전 보드 단독 검증"],
            ["향후 범위", "ADC 입력 보정, CAN transceiver 연결, BAMOCAR 응답 확인"],
        ],
        [2200, 7160],
    )

    add_heading(doc, "1. 목적", 1)
    add_body(
        doc,
        "본 보고서는 STM32L432KC 보드를 이용하여 APPS 기반 모터 제어 시스템의 기본 동작 구조를 사전에 검토하고, "
        "CAN 통신 및 APPS 타당성 검사 로직의 구현 가능성을 확인하기 위해 작성하였다.",
    )
    add_body(
        doc,
        "최종 목표는 APPS 센서의 두 아날로그 신호를 STM32에서 읽고, 두 신호의 타당성을 검사한 뒤 정상 상태에서만 "
        "BAMOCAR 모터 컨트롤러로 CAN torque command를 전송하는 것이다.",
    )

    add_heading(doc, "2. 시스템 구성", 1)
    add_body(doc, "전체 제어 흐름은 다음과 같이 구성된다.")
    add_code_block(
        doc,
        [
            "APPS Sensor",
            "-> STM32 ADC",
            "-> APPS Plausibility Check",
            "-> Torque Command Calculation",
            "-> CAN Frame Generation",
            "-> BAMOCAR Motor Controller",
        ],
    )
    add_body(
        doc,
        "현재 단계에서는 실제 APPS 센서와 BAMOCAR가 모두 연결된 상태는 아니므로, STM32 보드 단독으로 확인 가능한 "
        "CLI 기반 사전 검증을 중심으로 진행하였다.",
    )

    add_heading(doc, "3. CAN 통신 구조", 1)
    add_body(doc, "BAMOCAR와의 통신은 Classic CAN 기준으로 설계하였다.")
    add_table(
        doc,
        ["구분", "설정값", "설명"],
        [
            ["STM32 -> BAMOCAR", "0x201", "BAMOCAR가 수신하는 명령 CAN ID"],
            ["BAMOCAR -> STM32", "0x181", "STM32가 수신하는 응답 CAN ID"],
            ["Baudrate", "500 kbit/s", "초기 검증 기준 통신 속도"],
            ["Frame type", "Classic CAN", "Standard ID 기반 데이터 프레임"],
        ],
        [2400, 2200, 4760],
    )
    add_body(doc, "BAMOCAR 상태 요청 프레임은 다음과 같이 구성된다.")
    add_table(
        doc,
        ["필드", "값", "의미"],
        [
            ["ID", "0x201", "STM32에서 BAMOCAR로 전송"],
            ["DLC", "3", "데이터 3 byte"],
            ["DATA", "3D 40 00", "KERN_STATUS 1회 읽기 요청"],
        ],
        [1800, 2500, 5060],
    )
    add_body(doc, "Torque command 프레임은 다음 형식을 사용한다.")
    add_table(
        doc,
        ["필드", "값", "의미"],
        [
            ["ID", "0x201", "STM32에서 BAMOCAR로 전송"],
            ["DLC", "3", "데이터 3 byte"],
            ["DATA", "90 LL HH", "0x90 torque command register와 little-endian torque 값"],
        ],
        [1800, 2500, 5060],
    )

    add_heading(doc, "4. APPS 타당성 검사 로직", 1)
    add_body(
        doc,
        "APPS는 두 개의 독립된 센서 신호를 사용한다. 각 신호는 ADC raw 값으로 입력되며, 이를 pedal position 값으로 변환한다.",
    )
    add_number(doc, "raw_ch1, raw_ch2 입력")
    add_number(doc, "각 채널을 0~1000 per-mille 값으로 변환")
    add_number(doc, "두 채널의 pedal 값 차이 계산")
    add_number(doc, "차이가 10%를 초과하면 fault 판단")
    add_number(doc, "fault 발생 시 torque command를 0으로 제한")
    add_number(doc, "정상 상태에서는 pedal 값에 비례하여 torque command 생성")
    add_body(doc, "Fault 상태는 latch 구조로 유지되며, CLI 명령을 통해 clear할 수 있도록 구성하였다.")

    add_heading(doc, "5. 사전 검증 내용", 1)
    add_body(doc, "STM32 보드 단독 상태에서 CLI 기반으로 APPS 로직을 사전 검증할 수 있도록 테스트 명령을 구성하였다.")
    add_table(
        doc,
        ["CLI 명령", "용도"],
        [
            ["apps test raw_ch1 raw_ch2", "두 APPS raw 입력값에 대한 pedal %, fault 상태, torque command 확인"],
            ["apps torque pedal_per_mille", "pedal per-mille 값에 대한 BAMOCAR torque frame 생성 확인"],
            ["apps clear", "latched fault 해제"],
        ],
        [3600, 5760],
    )
    add_heading(doc, "5.1 테스트 예시", 2)
    add_table(
        doc,
        ["명령", "기대 결과"],
        [
            ["apps test 1000 1000", "두 채널 값이 동일하므로 정상 상태로 판단되고 pedal 값과 torque command 계산"],
            ["apps test 1000 2500", "두 채널 차이가 커서 DIFF_FAULT 및 FAULT_LATCHED 발생, torque command 0 제한"],
            ["apps clear", "latched fault 해제"],
            ["apps torque 500", "pedal 50.0%에 해당하는 BAMOCAR torque command frame 확인"],
        ],
        [3600, 5760],
    )

    add_heading(doc, "6. UART/CLI 초기화 문제 및 해결", 1)
    add_body(
        doc,
        "사전 검증 중 CLI가 정상적으로 열리지 않는 문제가 발생하였다. 원인은 UART가 처음 open되기 전에 "
        "`HAL_UART_DeInit()`이 먼저 호출되면서, DMA handle이 아직 연결되지 않은 상태에서 DeInit 경로를 타는 것이었다.",
    )
    add_body(doc, "이를 해결하기 위해 UART가 이미 open된 경우에만 DeInit을 수행하도록 수정하였다.")
    add_code_block(doc, ["if (is_open[ch] == true)", "{", "  HAL_UART_DeInit(&huart2);", "}"])
    add_body(doc, "또한 DMA handle이 존재할 때만 DMA DeInit을 수행하도록 보호하였다.")
    add_code_block(doc, ["if (uartHandle->hdmarx != NULL)", "{", "  HAL_DMA_DeInit(uartHandle->hdmarx);", "}"])
    add_body(doc, "수정 후 CLI가 정상적으로 동작하는 것을 확인하였다.")

    add_heading(doc, "7. 현재 결과", 1)
    add_bullet(doc, "STM32 CLI 초기화 구조 정상화")
    add_bullet(doc, "CAN driver 및 CLI 명령 구조 확인")
    add_bullet(doc, "BAMOCAR CAN frame 생성 방식 정리")
    add_bullet(doc, "APPS plausibility 로직 구현")
    add_bullet(doc, "APPS raw 입력 기반 CLI 사전 검증 구조 구현")
    add_bullet(doc, "fault 발생 시 torque command 0 처리 구조 확인")
    add_body(
        doc,
        "이를 통해 실제 센서와 BAMOCAR가 연결되기 전에도 APPS 판정 로직과 CAN frame 생성 구조를 사전에 확인할 수 있는 기반을 마련하였다.",
    )

    add_heading(doc, "8. 향후 작업", 1)
    future_items = [
        "APPS Signal1, Signal2 실제 전압 범위 측정",
        "STM32 ADC 채널 연결 및 raw 값 확인",
        "APPS 보정값 확정",
        "100 ms 이상 fault 지속 조건 검증",
        "CAN transceiver 연결",
        "BAMOCAR 응답 ID 0x181 수신 확인",
        "torque 0%, 10% command 전송 확인",
        "fault 발생 시 torque 0 전송 확인",
        "BAMOCAR 상태값 및 fault 응답 확인",
    ]
    for item in future_items:
        add_bullet(doc, item)

    add_heading(doc, "9. 결론", 1)
    add_body(
        doc,
        "본 사전 검증을 통해 STM32 보드에서 APPS 기반 모터 제어를 위한 기본 소프트웨어 구조를 확인하였다. "
        "APPS raw 값 입력, pedal position 변환, plausibility fault 판단, torque command 생성, BAMOCAR CAN frame 생성 흐름을 "
        "CLI 기반으로 검증할 수 있도록 구현하였다.",
    )
    add_body(
        doc,
        "향후 실제 APPS 센서와 BAMOCAR를 연결하여 ADC 입력값 보정, CAN 응답 확인, fault 발생 시 torque 차단 동작을 "
        "실제 하드웨어 환경에서 검증할 예정이다.",
    )

    doc.core_properties.title = "APPS 기반 CAN 모터 제어 사전 검증 보고서"
    doc.core_properties.subject = "STM32L432KC APPS-CAN 사전 검증"
    doc.core_properties.author = "Codex"
    doc.save(OUT)


if __name__ == "__main__":
    build()
